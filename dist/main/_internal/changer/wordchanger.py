from docx import Document
import os

def replace_text_in_docx(docx_file, old_text, new_text):
    # Load the Word document
    try:
        doc = Document(docx_file)
    except:
        print(f"Word Datei {docx_file}")
        
    replaced = False

    # Iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        if old_text.lower() in paragraph.text.lower():
            # Iterate through the runs in the paragraph to maintain formatting
            for run in paragraph.runs:
                if old_text.lower() in run.text.lower():
                    # Replace the old text with the new text while maintaining formatting
                    run.text = run.text.replace(old_text, new_text)
                    replaced = True
                    break  # Stop iterating through runs once the text is replaced

            if replaced:
                break  # Stop iterating through paragraphs once the text is replaced

    # Iterate through all tables in the document
    if not replaced:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text.lower() in cell.text.lower():
                        # Iterate through the paragraphs in the cell to maintain formatting
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text.lower() in run.text.lower():
                                    # Replace the old text with the new text while maintaining formatting
                                    run.text = run.text.replace(old_text, new_text)
                                    replaced = True
                                    break  # Stop iterating through runs once the text is replaced

                            if replaced:
                                break  # Stop iterating through paragraphs once the text is replaced

            if replaced:
                break  # Stop iterating through tables once the text is replaced

    # Save the modified document if any replacement occurred
    if replaced:
        doc.save(docx_file)
        print("Document saved successfully.")
    else:
        print(f"Text '{old_text}' not found in the document.")

print(os.path.dirname(__file__))